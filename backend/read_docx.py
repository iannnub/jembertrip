import os
import sys

try:
    import docx
    print("python-docx is installed!")
except ImportError:
    print("python-docx is NOT installed. Installing it now...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx
    print("python-docx installed successfully!")

doc_path = r"D:\iann Kuliah\Semester 7\1. Artificial Intelegence (AI)\jembertrip\acc_fix_Proposal Tugas Akhir.docx"
if os.path.exists(doc_path):
    doc = docx.Document(doc_path)
    print("Word document loaded successfully!")
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    
    # Let's search for "5" or "klik" or "history"
    found = 0
    for idx, p in enumerate(doc.paragraphs):
        text = p.text.lower()
        if "5" in text and ("klik" in text or "history" in text or "riwayat" in text):
            print(f"[{idx}] {p.text[:150]}...")
            found += 1
    print(f"Found {found} matches in paragraphs.")
else:
    print(f"Document not found at: {doc_path}")
